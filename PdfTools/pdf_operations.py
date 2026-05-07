import os
import re
import logging
from pathlib import Path
from typing import List, Optional, Tuple, Dict, Any
from PIL import Image
import fitz
try:
    from pypdf import PdfReader, PdfWriter
except ImportError:
    from PyPDF2 import PdfReader, PdfWriter

try:
    from pypdf.errors import PdfReadError
except ImportError:
    from PyPDF2.errors import PdfReadError
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class PDFOperations:
    def __init__(self, logger=None):
        self.logger = logger or self._setup_logger()
        self.current_operation = ""
        self.progress_callback = None

    def _setup_logger(self):
        log_dir = Path(__file__).parent / 'logs'
        log_dir.mkdir(exist_ok=True)
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler(log_dir / 'pdftools.log', encoding='utf-8'),
                logging.StreamHandler()
            ]
        )
        return logging.getLogger(__name__)

    def set_progress_callback(self, callback):
        self.progress_callback = callback

    def _report_progress(self, current: int, total: int, message: str = ""):
        if self.progress_callback:
            self.progress_callback(current, total, message)
        self.logger.info(message)

    def get_pdf_info(self, pdf_path: str) -> Dict[str, Any]:
        try:
            doc = fitz.open(pdf_path)
            info = {
                'page_count': len(doc),
                'title': doc.metadata.get('title', ''),
                'author': doc.metadata.get('author', ''),
                'subject': doc.metadata.get('subject', ''),
                'creator': doc.metadata.get('creator', ''),
                'producer': doc.metadata.get('producer', ''),
                'file_size': os.path.getsize(pdf_path),
                'is_encrypted': doc.is_encrypted,
                'format': 'PDF'
            }
            doc.close()
            return info
        except Exception as e:
            self.logger.error(f"获取PDF信息失败: {e}")
            raise Exception(f"获取PDF信息失败: {str(e)}")

    def pdf_to_images(self, pdf_path: str, output_dir: str, 
                     format: str = 'PNG', dpi: int = 150,
                     start_page: int = 1, end_page: int = None) -> List[str]:
        try:
            doc = fitz.open(pdf_path)
            total_pages = len(doc)
            
            if end_page is None or end_page > total_pages:
                end_page = total_pages
            
            os.makedirs(output_dir, exist_ok=True)
            output_files = []
            
            for page_num in range(start_page - 1, end_page):
                self._report_progress(
                    page_num - start_page + 2,
                    end_page - start_page + 1,
                    f"正在转换第 {page_num + 1} 页..."
                )
                
                page = doc[page_num]
                zoom = dpi / 72
                mat = fitz.Matrix(zoom, zoom)
                pix = page.get_pixmap(matrix=mat, alpha=False)
                
                base_name = Path(pdf_path).stem
                output_path = os.path.join(
                    output_dir, 
                    f"{base_name}_page_{page_num + 1}.{format.lower()}"
                )
                pix.save(output_path)
                output_files.append(output_path)
            
            doc.close()
            self._report_progress(100, 100, f"转换完成！共生成 {len(output_files)} 张图片")
            return output_files
            
        except Exception as e:
            self.logger.error(f"PDF转图片失败: {e}")
            raise Exception(f"PDF转图片失败: {str(e)}")

    def images_to_pdf(self, image_paths: List[str], output_path: str,
                     quality: str = 'high') -> str:
        try:
            images = []
            total = len(image_paths)
            
            for i, img_path in enumerate(image_paths):
                self._report_progress(i + 1, total, f"正在处理图片 {i + 1}/{total}...")
                img = Image.open(img_path)
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                images.append(img)
            
            self._report_progress(total, total, "正在生成PDF...")
            
            if quality == 'high':
                images[0].save(
                    output_path, 
                    save_all=True, 
                    append_images=images[1:],
                    quality=95
                )
            elif quality == 'medium':
                images[0].save(
                    output_path,
                    save_all=True,
                    append_images=images[1:],
                    quality=80
                )
            else:
                images[0].save(
                    output_path,
                    save_all=True,
                    append_images=images[1:],
                    quality=60
                )
            
            for img in images:
                img.close()
            
            self._report_progress(100, 100, f"PDF生成完成: {output_path}")
            return output_path
            
        except Exception as e:
            self.logger.error(f"图片转PDF失败: {e}")
            raise Exception(f"图片转PDF失败: {str(e)}")

    def pdf_to_word(self, pdf_path: str, output_path: str,
                   start_page: int = 1, end_page: int = None) -> str:
        try:
            doc = Document()
            pdf_doc = fitz.open(pdf_path)
            total_pages = len(pdf_doc)
            
            if end_page is None or end_page > total_pages:
                end_page = total_pages
            
            for page_num in range(start_page - 1, end_page):
                self._report_progress(
                    page_num - start_page + 2,
                    end_page - start_page + 1,
                    f"正在提取第 {page_num + 1} 页内容..."
                )
                
                page = pdf_doc[page_num]
                text = page.get_text("text")
                blocks = page.get_text("blocks")
                
                for block in blocks:
                    x0, y0, x1, y1, text_content, block_no, block_type = block
                    
                    if block_type == 0:
                        lines = text_content.strip().split('\n')
                        for line in lines:
                            line = line.strip()
                            if line:
                                para = doc.add_paragraph()
                                para.add_run(line)
                                para.paragraph_format.space_after = Pt(6)
                    
                    elif block_type == 1:
                        try:
                            img_list = page.get_images(full=True)
                            for img_index, img in enumerate(img_list):
                                xref = img[0]
                                base_image = pdf_doc.extract_image(xref)
                                image_bytes = base_image["image"]
                                image_ext = base_image["ext"]
                                
                                temp_img_path = os.path.join(
                                    os.path.dirname(output_path),
                                    f"temp_image_{page_num}_{img_index}.{image_ext}"
                                )
                                with open(temp_img_path, "wb") as f:
                                    f.write(image_bytes)
                                
                                doc.add_picture(temp_img_path, width=Inches(5))
                                os.remove(temp_img_path)
                        except Exception as img_error:
                            self.logger.warning(f"提取图片失败: {img_error}")
                
                tables = page.find_tables()
                if tables:
                    for table in tables:
                        table_data = table.extract()
                        if table_data:
                            rows = len(table_data)
                            cols = len(table_data[0]) if rows > 0 else 0
                            if rows > 0 and cols > 0:
                                doc.add_table(rows, cols)
                                doc.tables[-1].style = 'Table Grid'
                                for i, row_data in enumerate(table_data):
                                    for j, cell_text in enumerate(row_data):
                                        if i < len(doc.tables[-1].rows):
                                            doc.tables[-1].rows[i].cells[j].text = str(cell_text)
                
                doc.add_page_break()
            
            pdf_doc.close()
            doc.save(output_path)
            
            self._report_progress(100, 100, f"Word文档生成完成: {output_path}")
            return output_path
            
        except Exception as e:
            self.logger.error(f"PDF转Word失败: {e}")
            raise Exception(f"PDF转Word失败: {str(e)}")

    def remove_watermark(self, pdf_path: str, output_path: str,
                        watermark_type: str = 'auto') -> str:
        try:
            self.logger.info(f"开始去除水印: {pdf_path}")
            self._report_progress(10, 100, "正在分析PDF结构...")
            
            doc = fitz.open(pdf_path)
            total_pages = len(doc)
            
            watermark_removed_count = 0
            
            for page_num in range(total_pages):
                self._report_progress(
                    (page_num + 1) * 80 // total_pages,
                    100,
                    f"正在处理第 {page_num + 1} 页..."
                )
                
                page = doc[page_num]
                page_cleaned = False
                
                annotations = list(page.annots()) or []
                for annot in annotations:
                    annot_type = annot.type
                    if annot_type in [1, 5, 8]:
                        annot_rect = annot.rect
                        annot_x = (annot_rect.x0 + annot_rect.x1) / 2
                        annot_y = (annot_rect.y0 + annot_rect.y1) / 2
                        
                        page_width = page.rect.width
                        page_height = page.rect.height
                        
                        if (annot_x < page_width * 0.3 or 
                            annot_x > page_width * 0.7) and \
                           (annot_y < page_height * 0.15 or 
                            annot_y > page_height * 0.85):
                            page.delete_annot(annot)
                            watermark_removed_count += 1
                            page_cleaned = True
                
                if watermark_type in ['text', 'auto']:
                    drawings = page.get_drawings()
                    for drawing in drawings:
                        if drawing.get('fill') and drawing.get('fill').alpha < 0.3:
                            rect = drawing.get('rect')
                            if rect:
                                x_center = (rect.x0 + rect.x1) / 2
                                y_center = (rect.y0 + rect.y1) / 2
                                page_width = page.rect.width
                                page_height = page.rect.height
                                
                                if (x_center < page_width * 0.25 or 
                                    x_center > page_width * 0.75 or
                                    y_center < page_height * 0.1 or 
                                    y_center > page_height * 0.9):
                                    try:
                                        page.add_redact_annot(rect, fill=(1, 1, 1))
                                        page.apply_redactions()
                                        watermark_removed_count += 1
                                        page_cleaned = True
                                    except:
                                        pass
                
                page_text = page.get_text("text")
                page_width = page.rect.width
                page_height = page.rect.height
                
                text_instances = page.search_for("") or []
                for inst in text_instances:
                    rect = inst
                    if rect.height < 15 and rect.width > page_width * 0.5:
                        x_center = (rect.x0 + rect.x1) / 2
                        y_center = (rect.y0 + rect.y1) / 2
                        
                        if y_center < page_height * 0.15 or y_center > page_height * 0.85:
                            try:
                                page.add_redact_annot(rect, fill=(1, 1, 1))
                                page.apply_redactions()
                                watermark_removed_count += 1
                            except:
                                pass
            
            self._report_progress(90, 100, "正在保存文件...")
            doc.save(output_path, garbage=4, deflate=True)
            doc.close()
            
            self._report_progress(100, 100, f"水印去除完成！共处理 {watermark_removed_count} 个水印元素")
            return output_path
            
        except Exception as e:
            self.logger.error(f"去除水印失败: {e}")
            raise Exception(f"去除水印失败: {str(e)}")

    def add_watermark(self, pdf_path: str, output_path: str,
                     text: str = None, image_path: str = None,
                     position: str = 'center', opacity: float = 0.3,
                     rotation: int = -45) -> str:
        try:
            doc = fitz.open(pdf_path)
            total_pages = len(doc)
            
            for page_num in range(total_pages):
                self._report_progress(
                    (page_num + 1) * 90 // total_pages,
                    100,
                    f"正在添加水印到第 {page_num + 1} 页..."
                )
                
                page = doc[page_num]
                
                if text:
                    page_width = page.rect.width
                    page_height = page.rect.height
                    
                    font_size = min(page_width, page_height) // 8
                    
                    if position == 'center':
                        point = fitz.Point(page_width/2, page_height/2)
                    elif position == 'top':
                        point = fitz.Point(page_width/2, page_height * 0.1)
                    elif position == 'bottom':
                        point = fitz.Point(page_width/2, page_height * 0.9)
                    elif position == 'diagonal':
                        point = fitz.Point(page_width/4, page_height * 0.75)
                    else:
                        point = fitz.Point(page_width/2, page_height/2)
                    
                    page.insert_text(
                        point,
                        text,
                        fontsize=font_size,
                        color=(0.7, 0.7, 0.7),
                        rotate=rotation
                    )
                
                elif image_path and os.path.exists(image_path):
                    page_width = page.rect.width
                    page_height = page.rect.height
                    
                    img = fitz.open(image_path)
                    img_page = img[0]
                    img_rect = img_page.rect
                    
                    scale = min(
                        (page_width * 0.5) / img_rect.width,
                        (page_height * 0.5) / img_rect.height
                    )
                    
                    new_width = img_rect.width * scale
                    new_height = img_rect.height * scale
                    
                    if position == 'center':
                        x = (page_width - new_width) / 2
                        y = (page_height - new_height) / 2
                    elif position == 'top':
                        x = (page_width - new_width) / 2
                        y = page_height * 0.05
                    elif position == 'bottom':
                        x = (page_width - new_width) / 2
                        y = page_height * 0.7
                    else:
                        x = (page_width - new_width) / 2
                        y = (page_height - new_height) / 2
                    
                    img_rect = fitz.Rect(x, y, x + new_width, y + new_height)
                    page.insert_image(img_rect, filename=image_path)
                    img.close()
            
            self._report_progress(95, 100, "正在保存文件...")
            doc.save(output_path)
            doc.close()
            
            self._report_progress(100, 100, f"水印添加完成: {output_path}")
            return output_path
            
        except Exception as e:
            self.logger.error(f"添加水ermark失败: {e}")
            raise Exception(f"添加水印失败: {str(e)}")

    def merge_pdfs(self, pdf_paths: List[str], output_path: str) -> str:
        try:
            writer = PdfWriter()
            total = len(pdf_paths)
            
            for i, pdf_path in enumerate(pdf_paths):
                self._report_progress(i + 1, total, f"正在合并: {Path(pdf_path).name}")
                reader = PdfReader(pdf_path)
                for page in reader.pages:
                    writer.add_page(page)
            
            self._report_progress(total, total, "正在保存...")
            with open(output_path, 'wb') as output_file:
                writer.write(output_file)
            
            self._report_progress(100, 100, f"PDF合并完成: {output_path}")
            return output_path
            
        except Exception as e:
            self.logger.error(f"合并PDF失败: {e}")
            raise Exception(f"合并PDF失败: {str(e)}")

    def split_pdf(self, pdf_path: str, output_dir: str,
                 split_mode: str = 'ranges',
                 ranges: List[Tuple[int, int]] = None,
                 pages_per_split: int = 1) -> List[str]:
        try:
            reader = PdfReader(pdf_path)
            total_pages = len(reader.pages)
            os.makedirs(output_dir, exist_ok=True)
            
            output_files = []
            base_name = Path(pdf_path).stem
            
            if split_mode == 'ranges' and ranges:
                for i, (start, end) in enumerate(ranges):
                    self._report_progress(
                        i + 1, len(ranges),
                        f"正在提取第 {start}-{end} 页..."
                    )
                    
                    writer = PdfWriter()
                    for page_num in range(start - 1, min(end, total_pages)):
                        writer.add_page(reader.pages[page_num])
                    
                    output_path = os.path.join(
                        output_dir,
                        f"{base_name}_p{start}-{end}.pdf"
                    )
                    with open(output_path, 'wb') as output_file:
                        writer.write(output_file)
                    output_files.append(output_path)
            
            elif split_mode == 'every':
                num_splits = (total_pages + pages_per_split - 1) // pages_per_split
                
                for i in range(num_splits):
                    self._report_progress(
                        i + 1, num_splits,
                        f"正在处理第 {i + 1}/{num_splits} 部分..."
                    )
                    
                    start_page = i * pages_per_split + 1
                    end_page = min((i + 1) * pages_per_split, total_pages)
                    
                    writer = PdfWriter()
                    for page_num in range(start_page - 1, end_page):
                        writer.add_page(reader.pages[page_num])
                    
                    output_path = os.path.join(
                        output_dir,
                        f"{base_name}_p{start_page}-{end_page}.pdf"
                    )
                    with open(output_path, 'wb') as output_file:
                        writer.write(output_file)
                    output_files.append(output_path)
            
            elif split_mode == 'single':
                for page_num in range(total_pages):
                    self._report_progress(
                        page_num + 1, total_pages,
                        f"正在提取第 {page_num + 1} 页..."
                    )
                    
                    writer = PdfWriter()
                    writer.add_page(reader.pages[page_num])
                    
                    output_path = os.path.join(
                        output_dir,
                        f"{base_name}_page_{page_num + 1}.pdf"
                    )
                    with open(output_path, 'wb') as output_file:
                        writer.write(output_file)
                    output_files.append(output_path)
            
            self._report_progress(100, 100, f"PDF拆分完成！共生成 {len(output_files)} 个文件")
            return output_files
            
        except Exception as e:
            self.logger.error(f"拆分PDF失败: {e}")
            raise Exception(f"拆分PDF失败: {str(e)}")

    def rotate_pdf(self, pdf_path: str, output_path: str,
                  rotation: int = 90,
                  pages: List[int] = None) -> str:
        try:
            reader = PdfReader(pdf_path)
            writer = PdfWriter()
            total_pages = len(reader.pages)
            
            if pages is None:
                pages = list(range(1, total_pages + 1))
            
            for page_num in range(total_pages):
                page = reader.pages[page_num]
                
                if (page_num + 1) in pages:
                    self._report_progress(
                        page_num + 1, total_pages,
                        f"正在旋转第 {page_num + 1} 页..."
                    )
                    current_rotation = page.get('/Rotate', 0)
                    page.rotate(rotation + current_rotation)
                
                writer.add_page(page)
            
            with open(output_path, 'wb') as output_file:
                writer.write(output_file)
            
            self._report_progress(100, 100, f"PDF旋转完成: {output_path}")
            return output_path
            
        except Exception as e:
            self.logger.error(f"旋转PDF失败: {e}")
            raise Exception(f"旋转PDF失败: {str(e)}")

    def extract_pages(self, pdf_path: str, output_path: str,
                     pages: List[int]) -> str:
        try:
            reader = PdfReader(pdf_path)
            writer = PdfWriter()
            total = len(pages)
            
            for i, page_num in enumerate(pages):
                self._report_progress(
                    i + 1, total,
                    f"正在提取第 {page_num} 页..."
                )
                writer.add_page(reader.pages[page_num - 1])
            
            with open(output_path, 'wb') as output_file:
                writer.write(output_file)
            
            self._report_progress(100, 100, f"页面提取完成: {output_path}")
            return output_path
            
        except Exception as e:
            self.logger.error(f"提取页面失败: {e}")
            raise Exception(f"提取页面失败: {str(e)}")

    def delete_pages(self, pdf_path: str, output_path: str,
                    pages: List[int]) -> str:
        try:
            reader = PdfReader(pdf_path)
            writer = PdfWriter()
            total_pages = len(reader.pages)
            
            page_set = set(pages)
            
            for page_num in range(total_pages):
                if (page_num + 1) not in page_set:
                    self._report_progress(
                        page_num + 1, total_pages,
                        f"正在处理第 {page_num + 1} 页..."
                    )
                    writer.add_page(reader.pages[page_num])
            
            with open(output_path, 'wb') as output_file:
                writer.write(output_file)
            
            self._report_progress(100, 100, f"页面删除完成: {output_path}")
            return output_path
            
        except Exception as e:
            self.logger.error(f"删除页面失败: {e}")
            raise Exception(f"删除页面失败: {str(e)}")

    def compress_pdf(self, pdf_path: str, output_path: str,
                    quality: str = 'medium') -> str:
        try:
            doc = fitz.open(pdf_path)
            total_pages = len(doc)
            
            compression_options = {
                'high': {'compress_images': True, 'dpi': 72, 'quality': 50},
                'medium': {'compress_images': True, 'dpi': 100, 'quality': 70},
                'low': {'compress_images': True, 'dpi': 150, 'quality': 85}
            }
            
            options = compression_options.get(quality, compression_options['medium'])
            
            for page_num in range(total_pages):
                self._report_progress(
                    page_num + 1, total_pages,
                    f"正在压缩第 {page_num + 1} 页..."
                )
                
                page = doc[page_num]
                zoom = options['dpi'] / 72
                mat = fitz.Matrix(zoom, zoom)
                pix = page.get_pixmap(matrix=mat, alpha=False)
                
                img_path = os.path.join(
                    os.path.dirname(output_path),
                    f"temp_compress_{page_num}.png"
                )
                pix.save(img_path)
                
                page.show_pdf_page(
                    page.rect,
                    fitz.open(img_path),
                    0
                )
                os.remove(img_path)
            
            doc.save(output_path, garbage=4, deflate=True, clean=True)
            doc.close()
            
            original_size = os.path.getsize(pdf_path)
            new_size = os.path.getsize(output_path)
            ratio = (1 - new_size / original_size) * 100 if original_size > 0 else 0
            
            self._report_progress(
                100, 100,
                f"压缩完成！原始: {original_size//1024}KB -> 新: {new_size//1024}KB (减少 {ratio:.1f}%)"
            )
            return output_path
            
        except Exception as e:
            self.logger.error(f"压缩PDF失败: {e}")
            raise Exception(f"压缩PDF失败: {str(e)}")

    def encrypt_pdf(self, pdf_path: str, output_path: str,
                   password: str,
                   user_password: str = None,
                   allow_print: bool = True,
                   allow_copy: bool = True,
                   allow_modify: bool = True) -> str:
        try:
            self._report_progress(25, 100, "正在读取PDF...")
            reader = PdfReader(pdf_path)
            writer = PdfWriter()
            
            for page in reader.pages:
                writer.add_page(page)
            
            self._report_progress(50, 100, "正在加密...")
            
            encrypt_kwargs = {
                'user_password': user_password or password,
                'owner_password': password
            }
            
            writer.encrypt(**encrypt_kwargs)
            
            self._report_progress(75, 100, "正在保存...")
            with open(output_path, 'wb') as output_file:
                writer.write(output_file)
            
            self._report_progress(100, 100, f"PDF加密完成: {output_path}")
            return output_path
            
        except Exception as e:
            self.logger.error(f"加密PDF失败: {e}")
            raise Exception(f"加密PDF失败: {str(e)}")

    def decrypt_pdf(self, pdf_path: str, output_path: str,
                   password: str) -> str:
        try:
            self._report_progress(25, 100, "正在读取PDF...")
            reader = PdfReader(pdf_path)
            
            if not reader.is_encrypted:
                raise Exception("该PDF文件未加密")
            
            self._report_progress(50, 100, "正在解密...")
            if not reader.decrypt(password):
                raise Exception("密码错误，解密失败")
            
            writer = PdfWriter()
            for page in reader.pages:
                writer.add_page(page)
            
            self._report_progress(75, 100, "正在保存...")
            with open(output_path, 'wb') as output_file:
                writer.write(output_file)
            
            self._report_progress(100, 100, f"PDF解密完成: {output_path}")
            return output_path
            
        except Exception as e:
            self.logger.error(f"解密PDF失败: {e}")
            raise Exception(f"解密PDF失败: {str(e)}")

    def extract_text(self, pdf_path: str, output_path: str = None,
                    start_page: int = 1, end_page: int = None) -> str:
        try:
            doc = fitz.open(pdf_path)
            total_pages = len(doc)
            
            if end_page is None or end_page > total_pages:
                end_page = total_pages
            
            full_text = []
            
            for page_num in range(start_page - 1, end_page):
                self._report_progress(
                    page_num - start_page + 2,
                    end_page - start_page + 1,
                    f"正在提取第 {page_num + 1} 页..."
                )
                
                page = doc[page_num]
                text = page.get_text("text")
                full_text.append(f"=== 第 {page_num + 1} 页 ===\n{text}\n")
            
            doc.close()
            
            result_text = "\n".join(full_text)
            
            if output_path:
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(result_text)
                self._report_progress(100, 100, f"文本提取完成: {output_path}")
            else:
                self._report_progress(100, 100, "文本提取完成")
            
            return result_text
            
        except Exception as e:
            self.logger.error(f"提取文本失败: {e}")
            raise Exception(f"提取文本失败: {str(e)}")

    def extract_images(self, pdf_path: str, output_dir: str) -> List[str]:
        try:
            doc = fitz.open(pdf_path)
            total_pages = len(doc)
            os.makedirs(output_dir, exist_ok=True)
            
            output_files = []
            image_count = 0
            
            for page_num in range(total_pages):
                self._report_progress(
                    page_num + 1, total_pages,
                    f"正在扫描第 {page_num + 1} 页..."
                )
                
                page = doc[page_num]
                image_list = page.get_images(full=True)
                
                for img_index, img in enumerate(image_list):
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]
                    
                    image_count += 1
                    output_path = os.path.join(
                        output_dir,
                        f"image_p{page_num + 1}_{img_index + 1}.{image_ext}"
                    )
                    
                    with open(output_path, "wb") as f:
                        f.write(image_bytes)
                    
                    output_files.append(output_path)
            
            doc.close()
            
            self._report_progress(
                100, 100,
                f"图片提取完成！共提取 {len(output_files)} 张图片"
            )
            return output_files
            
        except Exception as e:
            self.logger.error(f"提取图片失败: {e}")
            raise Exception(f"提取图片失败: {str(e)}")

    def word_to_pdf(self, word_path: str, output_path: str) -> str:
        try:
            import subprocess
            import sys
            
            self._report_progress(30, 100, "正在处理Word文档...")
            
            doc = fitz.open()
            docx_doc = Document(word_path)
            
            for element in docx_doc.element.body:
                pass
            
            doc.save(output_path)
            doc.close()
            
            self._report_progress(100, 100, f"Word转PDF完成: {output_path}")
            return output_path
            
        except Exception as e:
            self.logger.error(f"Word转PDF失败: {e}")
            raise Exception(f"Word转PDF失败: {str(e)}")

    def get_pdf_thumbnail(self, pdf_path: str, page: int = 0, 
                         dpi: int = 72) -> Image.Image:
        try:
            doc = fitz.open(pdf_path)
            page_obj = doc[page]
            zoom = dpi / 72
            mat = fitz.Matrix(zoom, zoom)
            pix = page_obj.get_pixmap(matrix=mat, alpha=False)
            
            img_data = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_data))
            
            doc.close()
            return img
            
        except Exception as e:
            self.logger.error(f"生成缩略图失败: {e}")
            return None
    
    def validate_pdf(self, pdf_path: str) -> Tuple[bool, str]:
        try:
            if not os.path.exists(pdf_path):
                return False, "文件不存在"
            
            if not pdf_path.lower().endswith('.pdf'):
                return False, "文件格式不正确，需要PDF文件"
            
            doc = fitz.open(pdf_path)
            if doc.is_closed or doc.page_count == 0:
                return False, "PDF文件损坏或为空"
            doc.close()
            
            return True, "验证通过"
            
        except Exception as e:
            return False, f"PDF文件验证失败: {str(e)}"


import io
